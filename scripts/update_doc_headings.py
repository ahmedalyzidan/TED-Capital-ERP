import sys
import re
from pathlib import Path
from collections import defaultdict

try:
    import pandas as pd
except Exception as e:
    print("Missing dependency: pandas. Install with 'pip install pandas openpyxl python-docx'")
    print(e)
    sys.exit(1)

try:
    import docx
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
except Exception as e:
    print("Missing dependency: python-docx. Install with 'pip install python-docx'")
    print(e)
    sys.exit(1)


EXCEL_PATH = Path(r"D:\z\Digital onboarding\AAIB_Digital_Onboarding_Requirements_Comparison_changes.xlsx")
DOC_PATH = Path(r"D:\z\Digital onboarding\AAIB_Business_Digital_Onboarding_RFP_VF 18-05-2026_Backup.docx")
UPDATED_DOC_PATH = Path(r"D:\z\Digital onboarding\AAIB_Business_Digital_Onboarding_RFP_VF 18-05-2026_updated.docx")
CHANGES_DOC_PATH = Path(r"D:\z\Digital onboarding\AAIB_Business_Digital_Onboarding_RFP_VF 18-05-2026_changes.docx")


def insert_paragraph_after(paragraph, text=None, style=None, italic=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        if italic:
            run.italic = True
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def build_mapping(excel_path: Path):
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")
    # Read with pandas
    df = pd.read_excel(excel_path)
    # Expected columns: 'requirement_id' and 'new_value' (based on the user's description)
    # Normalize column names
    cols = {c.lower(): c for c in df.columns}
    # Find requirement id column
    req_col = None
    new_col = None
    for k, v in cols.items():
        if 'requirement' in k and 'id' in k:
            req_col = v
        if k in ('new_value', 'new value', 'matched_heading') or 'new' in k and 'value' in k:
            new_col = v
    # fallback guesses
    if req_col is None:
        for k, v in cols.items():
            if 'id' in k:
                req_col = v
                break
    if new_col is None:
        for k, v in cols.items():
            if 'new' in k or 'matched' in k:
                new_col = v
                break
    # additional fallbacks for requirement id-like columns
    if req_col is None:
        for name in ('row', 'sheet', 'id', 'requirement'):
            for k, v in cols.items():
                if name in k:
                    req_col = v
                    break
            if req_col is not None:
                break
    # final fallback: pick the first column that's not the new_col
    if req_col is None:
        for v in df.columns:
            if v != new_col:
                req_col = v
                break
    if req_col is None or new_col is None:
        # Try common explicit names
        if 'requirement_id' in df.columns and 'new_value' in df.columns:
            req_col = 'requirement_id'
            new_col = 'new_value'
        else:
            raise ValueError(f"Could not find 'requirement_id' or 'new_value' columns in Excel. Found: {list(df.columns)}")

    mapping = defaultdict(list)
    for _, row in df.iterrows():
        req = row[req_col]
        newv = row[new_col]
        if pd.isna(newv):
            continue
        newv_str = str(newv).strip()
        if newv_str == '':
            continue
        mapping[newv_str].append(str(req))
    return mapping


def is_heading_paragraph(p):
    try:
        style_name = p.style.name
    except Exception:
        style_name = ''
    return bool(style_name and isinstance(style_name, str) and style_name.startswith('Heading'))


def main():
    try:
        mapping = build_mapping(EXCEL_PATH)
    except Exception as e:
        print(f"Error reading Excel mapping: {e}")
        mapping = {}

    try:
        if not DOC_PATH.exists():
            raise FileNotFoundError(f"Word file not found: {DOC_PATH}")
        doc = Document(DOC_PATH)
    except Exception as e:
        print(f"Error opening Word document: {e}")
        sys.exit(1)

    # Collect heading paragraphs in order
    headings = []  # list of (index_in_doc_paragraphs, paragraph)
    for idx, p in enumerate(doc.paragraphs):
        if is_heading_paragraph(p):
            headings.append((idx, p))

    rename_records = []  # (old, new)
    annotations_inserted = 0

    # Process renames: for each heading, if contains 'Loan' and within next 3 headings one contains 'Digital'
    for i, (para_idx, para) in enumerate(headings):
        text = para.text
        if not text:
            continue
        if re.search(r'\bLoan\b', text, flags=re.I):
            # look ahead up to next 3 headings
            lookahead = headings[i+1:i+4]
            found_digital = False
            for _, p2 in lookahead:
                if re.search(r'Digital', p2.text, flags=re.I):
                    found_digital = True
                    break
            if found_digital:
                new_text = re.sub(r'\bLoan\b', 'Digital Business Onboarding', text, flags=re.I)
                if new_text != text:
                    old = text
                    para.text = new_text
                    rename_records.append((old, new_text))

    # After renames, we may need to rebuild headings list because texts changed; but original paragraph objects are same
    # Insert BRD alignment paragraphs where heading exact text equals a mapping key
    # We'll iterate through doc.paragraphs and when heading matches, insert
    for idx, p in list(enumerate(doc.paragraphs)):
        if is_heading_paragraph(p):
            key = p.text.strip()
            if key in mapping and mapping[key]:
                reqs = mapping[key]
                text = "BRD alignment: " + ", ".join(reqs)
                insert_paragraph_after(p, text=text, italic=True)
                annotations_inserted += 1

    # Save updated document
    try:
        doc.save( str(UPDATED_DOC_PATH) )
    except Exception as e:
        print(f"Error saving updated doc: {e}")
        sys.exit(1)

    # Create changes document listing heading changes
    try:
        changes_doc = Document()
        table = changes_doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'old_heading'
        hdr_cells[1].text = 'new_heading'
        hdr_cells[2].text = 'matched_requirements'
        for old, new in rename_records:
            reqs = mapping.get(new, [])
            row_cells = table.add_row().cells
            row_cells[0].text = old
            row_cells[1].text = new
            row_cells[2].text = ", ".join(reqs)
        changes_doc.save(str(CHANGES_DOC_PATH))
    except Exception as e:
        print(f"Error creating changes doc: {e}")
        sys.exit(1)

    # Print concise summary
    print(f"Heading renames: {len(rename_records)}")
    print(f"Annotations inserted: {annotations_inserted}")
    print(f"Updated document saved to: {UPDATED_DOC_PATH}")
    print(f"Changes document saved to: {CHANGES_DOC_PATH}")


if __name__ == '__main__':
    main()
