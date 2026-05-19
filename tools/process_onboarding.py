import os
import re
import sys
import traceback
from difflib import SequenceMatcher
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt


EXCEL_PATH = r"D:\z\Digital onboarding\AAIB_Digital_Onboarding_Requirements_Comparison.xlsx"
EXCEL_UPDATED = r"D:\z\Digital onboarding\AAIB_Digital_Onboarding_Requirements_Comparison_updated.xlsx"
EXCEL_CHANGES = r"D:\z\Digital onboarding\AAIB_Digital_Onboarding_Requirements_Comparison_changes.xlsx"

WORD_PATH = r"D:\z\Digital onboarding\AAIB_Business_Digital_Onboarding_RFP_VF 18-05-2026_Backup.docx"
WORD_UPDATED = r"D:\z\Digital onboarding\AAIB_Business_Digital_Onboarding_RFP_VF 18-05-2026_updated.docx"
WORD_CHANGES = r"D:\z\Digital onboarding\AAIB_Business_Digital_Onboarding_RFP_VF 18-05-2026_changes.docx"


def normalize_key(s: str) -> str:
    if s is None:
        return ""
    s = str(s).lower()
    s = re.sub(r"[^a-z0-9]", "", s)
    return s


def find_column(df, target_name):
    # Exact match trimmed case-insensitive
    for col in df.columns:
        if str(col).strip().lower() == target_name.strip().lower():
            return col
    # Try normalized matching
    target_norm = normalize_key(target_name)
    for col in df.columns:
        if normalize_key(col) == target_norm:
            return col
    return None


def extract_headings(doc: Document):
    headings = []
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        style_name = ""
        try:
            style_name = p.style.name
        except Exception:
            style_name = ""
        if isinstance(style_name, str) and style_name.lower().startswith("heading"):
            # extract level
            m = re.search(r"(\d+)", style_name)
            level = int(m.group(1)) if m else 1
            if level < 1:
                level = 1
            if level > 4:
                level = min(level, 4)
            # collect context: subsequent paragraphs until next heading
            context_texts = []
            for j in range(i + 1, len(paragraphs)):
                try:
                    sname = paragraphs[j].style.name
                except Exception:
                    sname = ""
                if isinstance(sname, str) and sname.lower().startswith("heading"):
                    break
                if paragraphs[j].text and paragraphs[j].text.strip():
                    context_texts.append(paragraphs[j].text.strip())
            context = " \n ".join(context_texts)
            headings.append({"index": len(headings), "level": level, "text": p.text.strip(), "context": context, "paragraph": p})
    return headings


def insert_paragraph_after(paragraph, text, italic=True, smaller_font_pt=9):
    # Insert new paragraph after given paragraph
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    from docx.text.paragraph import Paragraph

    new_par = Paragraph(new_p, paragraph._parent)
    run = new_par.add_run(text)
    if italic:
        run.italic = True
    try:
        run.font.size = Pt(smaller_font_pt)
    except Exception:
        pass
    return new_par


def main():
    changes = []
    docx_changes = []
    sheets_processed = 0
    rows_updated = 0

    # Load Word
    try:
        doc = Document(WORD_PATH)
    except Exception as e:
        print(f"Error loading Word file: {WORD_PATH}: {e}")
        doc = None

    headings = []
    if doc:
        headings = extract_headings(doc)

    # Load Excel
    try:
        xls = pd.read_excel(EXCEL_PATH, sheet_name=None, header=2, engine='openpyxl')
    except FileNotFoundError:
        print(f"Excel file not found: {EXCEL_PATH}")
        xls = {}
    except Exception as e:
        print(f"Error loading Excel: {e}")
        xls = {}

    heading_texts = [h['text'] + "\n" + h['context'] for h in headings]

    # Map heading index to matched requirement IDs
    heading_to_reqs = {h['index']: [] for h in headings}

    for sheet_name, df in xls.items():
        sheets_processed += 1
        try:
            if df.empty:
                continue
        except Exception:
            continue

        # find columns
        col_aligned = find_column(df, 'Aligned BRD Description')
        col_req_id = find_column(df, 'Requirement ID')
        col_old = find_column(df, 'Old RFP Description')

        if col_aligned is None:
            print(f"Sheet '{sheet_name}': 'Aligned BRD Description' not found; trying normalized names or skipping sheet.")
        # if still none, we'll skip updating aligned col but can still match using old desc

        # Prepare text fields safely
        def get_cell_text(row, colname):
            try:
                val = row.get(colname, "") if isinstance(row, dict) or hasattr(row, 'get') else row[colname]
            except Exception:
                try:
                    val = row[colname]
                except Exception:
                    val = ""
            return "" if pd.isna(val) else str(val)

        # iterate rows
        for idx, row in df.iterrows():
            try:
                req_id = get_cell_text(row, col_req_id) if col_req_id is not None else ""
                old_desc = get_cell_text(row, col_old) if col_old is not None else ""
                aligned_desc = get_cell_text(row, col_aligned) if col_aligned is not None else ""

                text_for_search = old_desc.strip() or aligned_desc.strip()
                check_text = (req_id + " " + old_desc + " " + aligned_desc).lower()
                keywords = ['digital', 'onboard', 'business digital']
                if not any(k in check_text for k in keywords):
                    continue

                # find best heading
                best_score = 0.0
                best_hidx = None
                for hidx, htxt in enumerate(heading_texts):
                    score = SequenceMatcher(None, text_for_search, htxt).ratio() if text_for_search and htxt else 0.0
                    if score > best_score:
                        best_score = score
                        best_hidx = hidx

                old_value = aligned_desc if col_aligned is not None else ""
                new_value = old_value
                matched_heading = ""
                if best_score >= 0.2 and best_hidx is not None:
                    new_value = headings[best_hidx]['text']
                    matched_heading = new_value
                    heading_to_reqs[best_hidx].append(req_id)

                if new_value != old_value:
                    rows_updated += 1
                    # save into dataframe
                    if col_aligned is not None:
                        df.at[idx, col_aligned] = new_value

                    # compute excel row number (header is row 3)
                    header_row_num = 3
                    row_number = header_row_num + 1 + int(idx)

                    changes.append({
                        'sheet': sheet_name,
                        'row_number': row_number,
                        'requirement_id': req_id,
                        'column_name': 'Aligned BRD Description',
                        'old_value': old_value,
                        'new_value': new_value,
                        'match_score': best_score,
                        'matched_heading': matched_heading,
                    })
            except Exception:
                print(f"Error processing sheet {sheet_name} row {idx}:\n" + traceback.format_exc())

        # write back modified df to xls dict
        xls[sheet_name] = df

    # Save updated excel workbook
    try:
        with pd.ExcelWriter(EXCEL_UPDATED, engine='openpyxl') as writer:
            for sheet_name, df in xls.items():
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        print(f"Wrote updated Excel to: {EXCEL_UPDATED}")
    except Exception as e:
        print(f"Error saving updated Excel: {e}\n" + traceback.format_exc())

    # Save changes list to excel
    try:
        if changes:
            changes_df = pd.DataFrame(changes)
            changes_df.to_excel(EXCEL_CHANGES, index=False)
            print(f"Wrote Excel changes to: {EXCEL_CHANGES}")
        else:
            print("No changes to write to Excel changes file.")
    except Exception as e:
        print(f"Error writing Excel changes: {e}\n" + traceback.format_exc())

    # Modify Word doc
    if doc:
        # Loan -> Digital Business Onboarding replacements when next 3 headings contain Digital
        for i, h in enumerate(headings):
            try:
                if 'loan' in h['text'].lower():
                    found = False
                    for j in range(i + 1, min(i + 4, len(headings))):
                        if 'digital' in headings[j]['text'].lower():
                            found = True
                            break
                    if found:
                        old_heading_text = h['text']
                        new_heading_text = re.sub(r'(?i)loan', 'Digital Business Onboarding', old_heading_text)
                        # update paragraph text
                        p = h['paragraph']
                        p.text = new_heading_text
                        docx_changes.append({'old': old_heading_text, 'new': new_heading_text, 'matched_requirements': ''})

        # Insert BRD alignment paragraphs after headings that have matched reqs
        for hidx, reqs in heading_to_reqs.items():
            if not reqs:
                continue
            h = headings[hidx]
            p = h['paragraph']
            ids = ','.join([str(x) for x in reqs if x])
            insert_paragraph_after(p, f"BRD alignment: {ids}", italic=True, smaller_font_pt=9)

        try:
            doc.save(WORD_UPDATED)
            print(f"Saved updated Word to: {WORD_UPDATED}")
        except Exception as e:
            print(f"Error saving updated Word: {e}\n" + traceback.format_exc())

    # Save docx changes log
    try:
        log_doc = Document()
        if docx_changes:
            table = log_doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'old_heading'
            hdr_cells[1].text = 'new_heading'
            hdr_cells[2].text = 'matched_requirements'
            for r in docx_changes:
                row_cells = table.add_row().cells
                row_cells[0].text = r['old']
                row_cells[1].text = r['new']
                row_cells[2].text = r.get('matched_requirements', '')
        else:
            log_doc.add_paragraph('No heading replacements made.')
        log_doc.save(WORD_CHANGES)
        print(f"Saved Word changes log to: {WORD_CHANGES}")
    except Exception as e:
        print(f"Error saving Word changes log: {e}\n" + traceback.format_exc())

    # summary
    print("--- Summary ---")
    print(f"sheets_processed: {sheets_processed}")
    print(f"rows_updated: {rows_updated}")
    print(f"Excel updated path: {EXCEL_UPDATED}")
    print(f"Excel changes path: {EXCEL_CHANGES}")
    print(f"Word updated path: {WORD_UPDATED}")
    print(f"Word changes path: {WORD_CHANGES}")


if __name__ == '__main__':
    try:
        main()
    except Exception:
        print('Unhandled exception:')
        traceback.print_exc()
